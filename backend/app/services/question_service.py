"""
backend/app/services/question_service.py
Manages question bank — AI generation, file upload parsing, CRUD, analytics.
"""
import io
import json
import csv
from typing import List, Optional, Dict, Any
from datetime import datetime

from app.core.config import settings
from app.core.logger import get_logger
from app.db.interview_models import QuestionDocument, QuestionUploadDocument

logger = get_logger(__name__)


# ─── AI Question Generation ───────────────────────────────────────────────────

async def generate_questions_for_job(
    job_id: str,
    job_description: str,
    job_title: str,
    company: str = "",
    easy_count: int = None,
    medium_count: int = None,
    hard_count: int = None,
    existing_question_ids: Optional[List[str]] = None,
) -> List[str]:
    """
    Generate questions using AI and save to question bank.
    Returns list of saved QuestionDocument IDs.
    """
    easy = easy_count or settings.DEFAULT_EASY_COUNT
    medium = medium_count or settings.DEFAULT_MEDIUM_COUNT
    hard = hard_count or settings.DEFAULT_HARD_COUNT

    logger.info(f"Generating questions for: {job_title} (easy={easy}, medium={medium}, hard={hard})")

    from app.agents.interview_generator.agent import generate_questions
    questions_data = generate_questions(
        job_description=job_description,
        job_title=job_title,
        company=company,
        easy_count=easy,
        medium_count=medium,
        hard_count=hard,
    )

    saved_ids = []
    for q_data in questions_data:
        try:
            doc = QuestionDocument(
                job_id=job_id,
                skill=q_data.get("skill", ""),
                technology=q_data.get("technology", q_data.get("skill", "")),
                topic=q_data.get("topic", ""),
                type=q_data.get("type", "mcq"),
                difficulty=q_data.get("difficulty", "medium"),
                question_text=q_data.get("question_text", ""),
                options=q_data.get("options", []),
                correct_answers=q_data.get("correct_answers", []),
                explanation=q_data.get("explanation", ""),
                code_template=q_data.get("code_template", ""),
                language=q_data.get("language", "python"),
                test_cases=q_data.get("test_cases", []),
                time_limit_seconds=q_data.get("time_limit_seconds", settings.PER_QUESTION_TIME_SECONDS),
                blank_answer=q_data.get("blank_answer", ""),
                tags=q_data.get("tags", []),
                source="ai_generated",
            )
            if doc.question_text.strip():
                await doc.insert()
                saved_ids.append(str(doc.id))
        except Exception as e:
            logger.error(f"Failed to save question: {e}")

    logger.info(f"Saved {len(saved_ids)} AI-generated questions for {job_title}")
    return saved_ids


# ─── File Upload Parsing ──────────────────────────────────────────────────────

async def parse_questions_from_file(
    file_content: bytes,
    filename: str,
    job_id: Optional[str] = None,
) -> List[str]:
    """
    Parse questions from uploaded file (PDF, DOCX, XLSX, CSV).
    Returns list of saved QuestionDocument IDs.
    """
    upload_doc = QuestionUploadDocument(
        job_id=job_id,
        filename=filename,
        file_format=_detect_format(filename),
        status="processing",
    )
    await upload_doc.insert()

    try:
        ext = _detect_format(filename)
        raw_text = ""

        if ext == "pdf":
            raw_text = _parse_pdf(file_content)
        elif ext == "docx":
            raw_text = _parse_docx(file_content)
        elif ext in ("xlsx", "xls"):
            return await _parse_excel(file_content, job_id, upload_doc)
        elif ext == "csv":
            return await _parse_csv(file_content, job_id, upload_doc)
        else:
            raw_text = file_content.decode("utf-8", errors="ignore")

        # Use LLM to extract structured questions from raw text
        saved_ids = await _llm_extract_questions(raw_text, job_id)

        await upload_doc.set({
            "questions_extracted": len(saved_ids),
            "questions_saved": len(saved_ids),
            "status": "completed",
        })
        return saved_ids

    except Exception as e:
        logger.error(f"Failed to parse file {filename}: {e}")
        await upload_doc.set({"status": "failed", "error": str(e)})
        return []


def _detect_format(filename: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"
    return ext


def _parse_pdf(content: bytes) -> str:
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(content))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as e:
        logger.warning(f"PDF parse failed: {e}")
        return ""


def _parse_docx(content: bytes) -> str:
    try:
        import docx
        doc = docx.Document(io.BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception as e:
        logger.warning(f"DOCX parse failed: {e}")
        return ""


async def _parse_excel(content: bytes, job_id: Optional[str], upload_doc) -> List[str]:
    """Parse questions from Excel: expects columns Question, Type, Difficulty, OptionA-D, Answer, Explanation"""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(content))
        ws = wb.active
        headers = [str(c.value or "").strip().lower() for c in ws[1]]

        saved_ids = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            row_data = {headers[i]: str(v or "").strip() for i, v in enumerate(row) if i < len(headers)}
            q_doc = _excel_row_to_question(row_data, job_id)
            if q_doc and q_doc.question_text:
                await q_doc.insert()
                saved_ids.append(str(q_doc.id))

        await upload_doc.set({
            "questions_extracted": len(saved_ids),
            "questions_saved": len(saved_ids),
            "status": "completed",
        })
        return saved_ids
    except Exception as e:
        logger.error(f"Excel parse failed: {e}")
        await upload_doc.set({"status": "failed", "error": str(e)})
        return []


async def _parse_csv(content: bytes, job_id: Optional[str], upload_doc) -> List[str]:
    """Parse questions from CSV."""
    try:
        text = content.decode("utf-8", errors="ignore")
        reader = csv.DictReader(io.StringIO(text))
        saved_ids = []
        for row in reader:
            lower_row = {k.lower().strip(): v.strip() for k, v in row.items() if k}
            q_doc = _csv_row_to_question(lower_row, job_id)
            if q_doc and q_doc.question_text:
                await q_doc.insert()
                saved_ids.append(str(q_doc.id))

        await upload_doc.set({
            "questions_extracted": len(saved_ids),
            "questions_saved": len(saved_ids),
            "status": "completed",
        })
        return saved_ids
    except Exception as e:
        logger.error(f"CSV parse failed: {e}")
        await upload_doc.set({"status": "failed", "error": str(e)})
        return []


def _excel_row_to_question(row: dict, job_id: Optional[str]) -> Optional[QuestionDocument]:
    """Convert Excel row to QuestionDocument."""
    q_text = row.get("question", row.get("question_text", "")).strip()
    if not q_text:
        return None

    options = []
    correct = []
    for letter in ["a", "b", "c", "d", "e"]:
        opt_text = row.get(f"option{letter}", row.get(f"option {letter}", "")).strip()
        if opt_text:
            options.append({"id": letter, "text": opt_text})

    answer_raw = row.get("answer", row.get("correct_answer", "")).strip().lower()
    correct = [a.strip() for a in answer_raw.split(",") if a.strip()]

    return QuestionDocument(
        job_id=job_id,
        question_text=q_text,
        type=_normalize_type(row.get("type", "mcq")),
        difficulty=_normalize_difficulty(row.get("difficulty", "medium")),
        options=options,
        correct_answers=correct,
        explanation=row.get("explanation", ""),
        skill=row.get("skill", row.get("topic", "")),
        topic=row.get("topic", ""),
        source="uploaded",
    )


def _csv_row_to_question(row: dict, job_id: Optional[str]) -> Optional[QuestionDocument]:
    return _excel_row_to_question(row, job_id)


async def _llm_extract_questions(raw_text: str, job_id: Optional[str]) -> List[str]:
    """Use LLM to extract structured questions from unstructured text."""
    try:
        from langchain_groq import ChatGroq
        from pydantic import BaseModel, Field

        class ExtractedQuestion(BaseModel):
            question_text: str
            type: str = "mcq"
            difficulty: str = "medium"
            options: List[Dict[str, str]] = Field(default_factory=list)
            correct_answers: List[str] = Field(default_factory=list)
            explanation: str = ""
            skill: str = ""

        class ExtractedBatch(BaseModel):
            questions: List[ExtractedQuestion] = Field(default_factory=list)

        llm = ChatGroq(model="llama-3.3-70b-versatile", temperature=0, groq_api_key=settings.GROQ_API_KEY)

        prompt = (
            "Extract all interview questions from the following text. "
            "For each question identify: question_text, type (mcq/short_answer/coding/case_study), "
            "difficulty (easy/medium/hard), options (if MCQ), correct_answers, explanation, skill.\n\n"
            f"TEXT:\n{raw_text[:6000]}"  # Limit to avoid token overflow
        )

        result = llm.with_structured_output(ExtractedBatch).invoke(prompt)
        saved_ids = []
        for q in result.questions:
            if q.question_text.strip():
                doc = QuestionDocument(
                    job_id=job_id,
                    question_text=q.question_text,
                    type=_normalize_type(q.type),
                    difficulty=_normalize_difficulty(q.difficulty),
                    options=q.options,
                    correct_answers=q.correct_answers,
                    explanation=q.explanation,
                    skill=q.skill,
                    source="uploaded",
                )
                await doc.insert()
                saved_ids.append(str(doc.id))
        return saved_ids
    except Exception as e:
        logger.error(f"LLM question extraction failed: {e}")
        return []


# ─── Manual Questions ─────────────────────────────────────────────────────────

async def create_manual_question(data: dict) -> QuestionDocument:
    """Create a manually-written question."""
    doc = QuestionDocument(
        **{k: v for k, v in data.items() if k not in ("id", "_id")},
        source="manual",
    )
    await doc.insert()
    return doc


async def update_question(question_id: str, data: dict) -> Optional[QuestionDocument]:
    q = await QuestionDocument.get(question_id)
    if not q:
        return None
    data["updated_at"] = datetime.utcnow()
    await q.set({k: v for k, v in data.items() if k not in ("id", "_id")})
    return await QuestionDocument.get(question_id)


async def delete_question(question_id: str) -> bool:
    q = await QuestionDocument.get(question_id)
    if q:
        await q.delete()
        return True
    return False


# ─── Question Bank Queries ────────────────────────────────────────────────────

async def list_questions(
    job_id: Optional[str] = None,
    skill: Optional[str] = None,
    difficulty: Optional[str] = None,
    q_type: Optional[str] = None,
    limit: int = 100,
    skip: int = 0,
) -> List[QuestionDocument]:
    query = {}
    if job_id:
        query["job_id"] = job_id
    if skill:
        query["skill"] = {"$regex": skill, "$options": "i"}
    if difficulty:
        query["difficulty"] = difficulty
    if q_type:
        query["type"] = q_type
    return await QuestionDocument.find(query).skip(skip).limit(limit).to_list()


async def get_question_bank_analytics() -> Dict[str, Any]:
    """Analytics on question bank quality."""
    all_q = await QuestionDocument.find().to_list()
    total = len(all_q)
    by_type = {}
    by_difficulty = {}
    by_skill = {}
    hard_questions = []
    easy_questions = []

    for q in all_q:
        by_type[q.type] = by_type.get(q.type, 0) + 1
        by_difficulty[q.difficulty] = by_difficulty.get(q.difficulty, 0) + 1
        if q.skill:
            by_skill[q.skill] = by_skill.get(q.skill, 0) + 1
        if q.pass_rate < 0.25 and q.times_used > 5:
            hard_questions.append({"id": str(q.id), "text": q.question_text[:80], "pass_rate": q.pass_rate})
        if q.pass_rate > 0.9 and q.times_used > 5:
            easy_questions.append({"id": str(q.id), "text": q.question_text[:80], "pass_rate": q.pass_rate})

    return {
        "total": total,
        "by_type": by_type,
        "by_difficulty": by_difficulty,
        "by_skill": dict(sorted(by_skill.items(), key=lambda x: -x[1])[:10]),
        "hardest_questions": hard_questions[:5],
        "easiest_questions": easy_questions[:5],
        "avg_pass_rate": round(sum(q.pass_rate for q in all_q) / max(total, 1), 3),
    }


# ─── Hybrid Mode ─────────────────────────────────────────────────────────────

async def build_hybrid_question_set(
    job_id: str,
    job_description: str,
    job_title: str,
    company: str,
    manual_ids: List[str],
    ai_count: int = 20,
) -> List[str]:
    """
    Option D: Company provides some questions + AI generates the rest.
    Returns combined list of question IDs.
    """
    ai_ids = await generate_questions_for_job(
        job_id=job_id,
        job_description=job_description,
        job_title=job_title,
        company=company,
        easy_count=ai_count // 3,
        medium_count=ai_count // 2,
        hard_count=ai_count - (ai_count // 3) - (ai_count // 2),
    )
    # Merge manual + AI, manual first
    all_ids = manual_ids + [aid for aid in ai_ids if aid not in manual_ids]
    return all_ids


# ─── Helpers ─────────────────────────────────────────────────────────────────

def _normalize_type(raw: str) -> str:
    raw = raw.lower().replace(" ", "_").replace("-", "_")
    valid = {"mcq", "true_false", "fill_blank", "short_answer", "long_answer", "coding", "case_study", "sql", "system_design", "debugging", "output_prediction"}
    return raw if raw in valid else "mcq"


def _normalize_difficulty(raw: str) -> str:
    raw = raw.lower()
    if raw in ("easy", "simple", "basic"):     return "easy"
    if raw in ("hard", "difficult", "advanced"): return "hard"
    return "medium"
